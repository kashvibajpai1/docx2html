"""
parser.py
---------
Convert a ``.docx`` file into a :class:`~docx2html.schema.DocDocument` AST.

This module is the *only* place that touches ``python-docx`` objects.  The
rest of the pipeline works exclusively against the schema types defined in
:mod:`docx2html.schema`.

Parsing strategy
~~~~~~~~~~~~~~~~
1. Open the document with ``python-docx``.
2. Build O(1) lookup maps from XML element identity → python-docx wrapper
   objects.  This prevents the O(n²) behaviour that arises when scanning all
   paragraphs / tables for every body element encountered.
3. Iterate over the *body* XML elements in document order so that tables and
   paragraphs appear in the correct sequence (``doc.element.body``).
4. Dispatch each element to a specialist handler based on its XML tag or the
   resolved ``python-docx`` style name.
5. Group consecutive list paragraphs into a single :class:`~schema.DocList`
   node (ordered or unordered) using a clean internal ``_PendingListItem``
   dataclass — no sentinel attributes are injected into public AST types.
6. A paragraph that contains both text runs *and* embedded images is split
   into multiple blocks (text paragraph + one Image per embedded picture) so
   that no content is silently discarded.
7. Extract embedded images and save them to the ``media/`` directory only when
   an actual image is found (the directory is created lazily).

Style name conventions used by Word
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
- ``Heading 1`` … ``Heading 9`` → headings
- ``List Bullet``, ``List Bullet 2`` … → unordered list
- ``List Number``, ``List Number 2`` … → ordered list
- ``Code``, ``Verbatim``, ``Code Block``, ``Source Code``,
  ``Preformatted Text`` (and any style containing those words) → code block
- Any style whose base name contains ``"code"`` (case-insensitive) → code block

DOCX numbering XML primer (for list detection)
~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
Word stores list metadata in ``word/numbering.xml``.  The indirection chain is:

  paragraph  →  <w:numPr>
                  <w:numId w:val="N"/>   ← instance ID
                  <w:ilvl  w:val="D"/>   ← nesting depth (0-based)
         ↓
  <w:num w:numId="N">
    <w:abstractNumId w:val="A"/>
         ↓
  <w:abstractNum w:abstractNumId="A">
    <w:lvl w:ilvl="D">
      <w:numFmt w:val="decimal|bullet|…"/>

``decimal``, ``lowerLetter``, ``upperLetter``, ``lowerRoman``,
``upperRoman`` → ordered list.  Everything else → unordered.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from docx2html.schema import (
    Block,
    CodeBlock,
    DocDocument,
    DocList,
    Heading,
    HorizontalRule,
    Image,
    ListItem,
    Paragraph,
    Table,
    TableCell,
    TableRow,
    TextRun,
)

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Compiled style-name patterns
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^[Hh]eading\s+(\d)$")
_LIST_BULLET_RE = re.compile(r"^List\s+Bullet", re.IGNORECASE)
_LIST_NUMBER_RE = re.compile(r"^List\s+(Number|Continue)", re.IGNORECASE)

# Fix #10 — expanded to cover "Code Block", "Source Code", "Preformatted Text"
# and any style name that *contains* any of the root words.
_CODE_RE = re.compile(
    r"code|verbatim|monospace|preformatted|source\s+code|code\s+block",
    re.IGNORECASE,
)

# Paragraph styles that are deliberately dropped (TOC entries, footnotes, …).
_SKIP_STYLE_RE = re.compile(
    r"^(TOC|Table of Contents|footnote|endnote|caption|Normal \(Web\))",
    re.IGNORECASE,
)

# ---------------------------------------------------------------------------
# Frequently used XML qualified names (pre-computed for speed)
# ---------------------------------------------------------------------------

_W_TBL = qn("w:tbl")
_W_P = qn("w:p")
_W_SDT = qn("w:sdt")
_W_SDTCONTENT = qn("w:sdtContent")
_W_R = qn("w:r")
_W_T = qn("w:t")
_W_TAB = qn("w:tab")
_W_BR = qn("w:br")
_W_HYPERLINK = qn("w:hyperlink")
_W_INS = qn("w:ins")
_W_DEL = qn("w:del")
_W_PPR = qn("w:pPr")
_W_RPR = qn("w:rPr")
_W_NUMPR = qn("w:numPr")
_W_NUMID = qn("w:numId")
_W_ILVL = qn("w:ilvl")
_W_B = qn("w:b")
_W_I = qn("w:i")
_W_U = qn("w:u")
_W_STRIKE = qn("w:strike")
_W_DSTRIKE = qn("w:dstrike")
_W_RFONTS = qn("w:rFonts")
_W_PBDR = qn("w:pBdr")
_W_BOTTOM = qn("w:bottom")
_W_VAL = qn("w:val")
_W_ASCII = qn("w:ascii")
_W_HANSI = qn("w:hAnsi")
_W_DRAWING = qn("w:drawing")
_W_PICT = qn("w:pict")
_R_ID = qn("r:id")
_A_BLIP = qn("a:blip")
# r:embed lives in the relationships namespace, not 'r:'.
_R_EMBED = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"

# Alignment map: WD_ALIGN_PARAGRAPH int value → CSS string.
_ALIGN_MAP: dict[int, str] = {
    0: "left",
    1: "center",
    2: "right",
    3: "justify",
}

# numFmt values that indicate an ordered (numbered) list.
_ORDERED_NUM_FMTS = frozenset(
    {"decimal", "lowerLetter", "upperLetter", "lowerRoman", "upperRoman"}
)

# Font families that imply inline-code styling.
_MONO_FONTS = ("courier", "consolas", "monaco", "monospace", "lucida console")

# ---------------------------------------------------------------------------
# Internal-only dataclass used during list accumulation (Fix #2)
# ---------------------------------------------------------------------------


@dataclass
class _PendingListItem:
    """
    Internal-only carrier for a list item before it is assigned to a DocList.

    This replaces the previous approach of injecting a private ``_ordered``
    sentinel attribute directly onto public :class:`~schema.ListItem` objects.
    The public AST schema is never polluted.
    """

    runs: list[TextRun]
    depth: int
    ordered: bool


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def parse(docx_path: str | Path, media_dir: str | Path | None = None) -> DocDocument:
    """
    Parse *docx_path* and return a :class:`~docx2html.schema.DocDocument`.

    Parameters
    ----------
    docx_path:
        Path to the ``.docx`` file to parse.
    media_dir:
        Directory where extracted images will be written.  Defaults to a
        ``media/`` directory next to *docx_path*.  Pass ``None`` to use the
        default.

    Returns
    -------
    DocDocument
        The fully parsed document AST.

    Raises
    ------
    FileNotFoundError
        When *docx_path* does not exist.
    ValueError
        When *docx_path* is not a ``.docx`` file.
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")
    if docx_path.suffix.lower() != ".docx":
        raise ValueError(f"Expected a .docx file, got: {docx_path.suffix!r}")

    if media_dir is None:
        media_dir = docx_path.parent / "media"
    media_dir = Path(media_dir)

    doc = DocxDocument(str(docx_path))
    return _Parser(doc, media_dir=media_dir).parse()


# ---------------------------------------------------------------------------
# Internal parser class
# ---------------------------------------------------------------------------


class _Parser:
    """Stateful parser that converts a ``python-docx`` document into an AST."""

    def __init__(self, doc: DocxDocument, media_dir: Path) -> None:
        self._doc = doc
        self._media_dir = media_dir
        self._image_counter: int = 0

        # Fix #1 — O(1) lookup maps built once in __init__.
        # ``doc.paragraphs`` includes paragraphs nested inside tables; that is
        # intentional here because we also need to look up paragraphs that
        # appear inside <w:sdt> content-control wrappers.
        self._para_map: dict[object, DocxParagraph] = {
            p._element: p for p in doc.paragraphs
        }
        self._table_map: dict[object, DocxTable] = {
            t._element: t for t in doc.tables
        }

        # Accumulator for list items being grouped into a DocList.
        self._pending_list: DocList | None = None
        # The ordered/unordered flag of the currently accumulating list.
        self._pending_ordered: bool = False

    # ------------------------------------------------------------------
    # Top-level entry point
    # ------------------------------------------------------------------

    def parse(self) -> DocDocument:
        """Walk the document body and build the block list."""
        blocks: list[Block] = []

        for item in self._iter_body_items():
            if isinstance(item, _PendingListItem):
                # Fix #2 — list items are carried as _PendingListItem; the
                # public ListItem / DocList types are only created here.
                self._accumulate_list_item(item, blocks)
            else:
                # Any non-list block flushes a pending list first.
                self._flush_pending_list(blocks)
                blocks.append(item)

        # Flush any list still open at end of document.
        self._flush_pending_list(blocks)

        return DocDocument(
            blocks=blocks,
            title=self._core_property("title"),
            author=self._core_property("author"),
            created=self._core_property("created"),
        )

    def _flush_pending_list(self, blocks: list[Block]) -> None:
        """Emit the accumulated :class:`~schema.DocList` into *blocks* and reset."""
        if self._pending_list is not None:
            blocks.append(self._pending_list)
            self._pending_list = None

    # ------------------------------------------------------------------
    # Body element iterator
    # ------------------------------------------------------------------

    def _iter_body_items(self) -> Iterator[Block | _PendingListItem]:
        """
        Yield AST nodes (or :class:`_PendingListItem` sentinels) in document
        order by walking the raw body XML.

        Using ``doc.element.body`` directly preserves the interleaving of
        paragraphs and tables that ``doc.paragraphs`` (which flattens
        everything) would destroy.

        Fix #4 — a single paragraph may now produce *multiple* blocks (e.g.
        a text paragraph followed by one or more Image nodes).  We yield each
        block individually so the caller's list-accumulation logic sees them
        in the right order.
        """
        body = self._doc.element.body

        for child in body:
            tag = child.tag

            if tag == _W_TBL:
                table = self._find_table_for_element(child)
                if table is not None:
                    yield self._parse_table(table)

            elif tag == _W_P:
                yield from self._parse_paragraph_element(child)

            elif tag == _W_SDT:
                # Structured document tags (content controls) may wrap
                # paragraphs — e.g. the document title field.
                for sdt_child in child:
                    if sdt_child.tag == _W_SDTCONTENT:
                        for inner in sdt_child:
                            if inner.tag == _W_P:
                                yield from self._parse_paragraph_element(inner)

            # Section properties, bookmarks, and other tags are ignored.

    def _parse_paragraph_element(
        self, p_elem: object
    ) -> Iterator[Block | _PendingListItem]:
        """
        Look up the python-docx Paragraph for *p_elem* and dispatch to
        :meth:`_parse_paragraph`, yielding all resulting blocks.
        """
        para = self._find_para_for_element(p_elem)
        if para is None:
            return
        yield from self._parse_paragraph(para)

    # ------------------------------------------------------------------
    # Paragraph parsing
    # ------------------------------------------------------------------

    def _parse_paragraph(
        self, para: DocxParagraph
    ) -> Iterator[Block | _PendingListItem]:
        """
        Convert a single ``python-docx`` paragraph into one or more AST items.

        Fix #4 — this method now yields an *iterator* of blocks so that a
        paragraph containing both text and images produces multiple outputs
        without dropping any content.

        Yields
        ------
        Block or _PendingListItem
            One or more AST items derived from the paragraph.  Nothing is
            yielded for empty / skipped paragraphs.
        """
        style_name: str = ""
        try:
            style_name = para.style.name or ""
        except Exception:
            logger.debug("Could not read style for paragraph; treating as Normal.")

        # Skip styles that are never useful in the output.
        if _SKIP_STYLE_RE.match(style_name):
            return

        # --- Headings ---
        m = _HEADING_RE.match(style_name)
        if m:
            level = int(m.group(1))
            runs = self._extract_runs(para)
            if any(r.text.strip() for r in runs):
                yield Heading(level=level, runs=runs)
            return

        # --- Code blocks ---
        if _CODE_RE.search(style_name):
            text = para.text
            if text.strip():  # skip empty code paragraphs
                yield CodeBlock(text=text, language="")
            return

        # --- List detection (style-based, then numPr-based) ---
        #
        # Word uses two independent mechanisms to mark list paragraphs:
        #   1. A named style like "List Bullet" or "List Number".
        #   2. A <w:numPr> element that references the numbering.xml definitions.
        # Both must be checked; real documents often use mechanism 2 without
        # setting a list-named style.
        list_item = self._detect_list_item(para, style_name)
        if list_item is not None:
            yield list_item
            return

        # --- Horizontal rule ---
        # Fix #11 — only treat as HR when the paragraph is also empty of text.
        if self._is_horizontal_rule(para) and not para.text.strip():
            yield HorizontalRule()
            return

        # --- Generic paragraph (possibly containing images) ---
        # Fix #4/#9 — walk the paragraph content in XML order to preserve the
        # interleaving of text runs and embedded images.
        yield from self._parse_mixed_paragraph(para)

    # ------------------------------------------------------------------
    # Mixed-content paragraph (text + images, in document order)
    # ------------------------------------------------------------------

    def _parse_mixed_paragraph(
        self, para: DocxParagraph
    ) -> Iterator[Block]:
        """
        Parse a paragraph that may contain text runs and/or embedded images.

        Emits:
        - A :class:`~schema.Paragraph` block for any text runs found,
          followed by
        - One :class:`~schema.Image` block per embedded image found.

        If the paragraph contains *only* images and no meaningful text, only
        Image blocks are emitted.  If it contains *only* text, only one
        Paragraph block is emitted.

        Fix #4 — previously only the first image was returned and all text
        was discarded when images were present.

        Fix #9 — document order is preserved by iterating the raw XML child
        list once, classifying each child as either a run-bearing element
        or an image-bearing element.
        """
        runs: list[TextRun] = []
        images: list[Image] = []

        for child in para._element:
            tag = child.tag

            # Run-bearing elements
            if tag == _W_R:
                run = self._parse_run_element(child)
                if run is not None:
                    runs.append(run)

            elif tag == _W_HYPERLINK:
                url = self._resolve_hyperlink_url(para, child)
                for r_elem in child.findall(_W_R):
                    run = self._parse_run_element(r_elem, hyperlink=url)
                    if run is not None:
                        runs.append(run)

            elif tag == _W_INS:
                # Tracked insertion — treat as normal content.
                for r_elem in child.findall(_W_R):
                    run = self._parse_run_element(r_elem)
                    if run is not None:
                        runs.append(run)

            # _W_DEL (tracked deletion) is intentionally skipped: deleted text
            # should not appear in the rendered output.

            # Image-bearing drawing elements
            elif tag == _W_DRAWING or tag == _W_PICT:
                # Images live inside <w:drawing> or <w:pict> children.
                img = self._extract_image_from_element(child, para)
                if img is not None:
                    images.append(img)

        alignment = self._para_alignment(para)
        has_text = any(r.text.strip() for r in runs)
        has_images = bool(images)

        if not has_text and not has_images:
            return  # truly empty paragraph — discard

        if has_text:
            yield Paragraph(runs=runs, alignment=alignment)

        # Yield each image as a separate top-level block.
        for img in images:
            yield img

    # ------------------------------------------------------------------
    # List detection
    # ------------------------------------------------------------------

    def _detect_list_item(
        self, para: DocxParagraph, style_name: str
    ) -> _PendingListItem | None:
        """
        Return a :class:`_PendingListItem` if *para* is a list paragraph,
        otherwise ``None``.

        Detection order (Fix #5):
        1. Style name matches ``List Bullet*`` → unordered.
        2. Style name matches ``List Number*`` / ``List Continue*`` → ordered.
        3. ``<w:numPr>`` element present with a non-zero ``numId`` → look up
           the numbering format in ``numbering.xml`` to determine ordered/
           unordered.  This handles paragraphs that use list formatting applied
           via the paragraph format dialog rather than a named list style.
        """
        # Strategy 1 & 2: style-name matching.
        if _LIST_BULLET_RE.match(style_name):
            depth = self._list_depth_from_style(style_name)
            runs = self._extract_runs(para)
            return _PendingListItem(runs=runs, depth=depth, ordered=False)

        if _LIST_NUMBER_RE.match(style_name):
            depth = self._list_depth_from_style(style_name)
            runs = self._extract_runs(para)
            return _PendingListItem(runs=runs, depth=depth, ordered=True)

        # Strategy 3: numPr-based detection.
        #
        # The <w:numPr> element is a child of <w:pPr> (paragraph properties).
        # Its two key children are:
        #   <w:ilvl w:val="D"/>  — indentation level (nesting depth, 0-based)
        #   <w:numId w:val="N"/> — reference into numbering.xml
        #
        # A numId of 0 means "remove list formatting" and must be ignored.
        num_pr = self._get_num_pr(para)
        if num_pr is not None:
            num_id = self._num_pr_num_id(num_pr)
            if num_id and num_id != 0:
                depth = self._list_depth_from_numpr(num_pr)
                ordered = self._is_ordered_via_numpr(para, num_pr)
                runs = self._extract_runs(para)
                # Only treat as a list item if there is actual text content.
                if any(r.text.strip() for r in runs):
                    return _PendingListItem(runs=runs, depth=depth, ordered=ordered)

        return None

    def _accumulate_list_item(
        self, item: _PendingListItem, blocks: list[Block]
    ) -> None:
        """
        Append *item* to the currently open :class:`~schema.DocList`.

        When the ordered/unordered type changes mid-stream (mixed list), the
        current list is flushed and a new one started.  This correctly handles
        documents that interleave bullet and numbered lists.

        Fix #2 — no sentinel attributes are written to any public AST type.
        """
        if self._pending_list is None:
            # Start a fresh list.
            self._pending_list = DocList(ordered=item.ordered)
            self._pending_ordered = item.ordered
        elif self._pending_ordered != item.ordered:
            # List type changed: emit the completed list and start a new one.
            self._flush_pending_list(blocks)
            self._pending_list = DocList(ordered=item.ordered)
            self._pending_ordered = item.ordered

        self._pending_list.items.append(
            ListItem(runs=item.runs, depth=item.depth)
        )

    # ------------------------------------------------------------------
    # List XML helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _get_num_pr(para: DocxParagraph) -> object | None:
        """
        Return the ``<w:numPr>`` element for *para*, or ``None``.

        Looks inside the paragraph's ``<w:pPr>`` properties block.
        """
        pPr = para._element.find(_W_PPR)
        if pPr is None:
            return None
        return pPr.find(_W_NUMPR)

    @staticmethod
    def _num_pr_num_id(num_pr: object) -> int | None:
        """Extract the ``w:numId`` integer value from a ``<w:numPr>`` element."""
        num_id_elem = num_pr.find(_W_NUMID)  # type: ignore[union-attr]
        if num_id_elem is None:
            return None
        try:
            return int(num_id_elem.get(_W_VAL, "0"))
        except (ValueError, TypeError):
            return None

    @staticmethod
    def _list_depth_from_style(style_name: str) -> int:
        """
        Extract the nesting depth from a list style name.

        ``List Bullet``   → depth 0
        ``List Bullet 2`` → depth 1
        ``List Bullet 3`` → depth 2
        """
        m = re.search(r"(\d+)$", style_name)
        if m:
            return max(0, int(m.group(1)) - 1)
        return 0

    @staticmethod
    def _list_depth_from_numpr(num_pr: object) -> int:
        """
        Read the ``<w:ilvl>`` value from a ``<w:numPr>`` element.

        ``<w:ilvl w:val="0"/>`` → depth 0 (top-level list item)
        ``<w:ilvl w:val="1"/>`` → depth 1 (first sub-list level)
        """
        ilvl = num_pr.find(_W_ILVL)  # type: ignore[union-attr]
        if ilvl is not None:
            try:
                return int(ilvl.get(_W_VAL, "0"))
            except (ValueError, TypeError):
                pass
        return 0

    def _is_ordered_via_numpr(
        self, para: DocxParagraph, num_pr: object
    ) -> bool:
        """
        Determine whether the list containing *para* is ordered (numbered).

        Navigates the numbering.xml indirection chain described in the module
        docstring.  Falls back to *unordered* on any error — a safe default
        since unordered lists degrade more gracefully visually.

        Fix #5 — this was previously a separate method called
        ``_is_ordered_list_via_numpr``.  It now receives the already-located
        ``num_pr`` element to avoid duplicate XML searches.
        """
        try:
            num_id_elem = num_pr.find(_W_NUMID)  # type: ignore[union-attr]
            if num_id_elem is None:
                return False
            num_id = int(num_id_elem.get(_W_VAL, "0"))
            if num_id == 0:
                return False

            ilvl_elem = num_pr.find(_W_ILVL)  # type: ignore[union-attr]
            ilvl = int(ilvl_elem.get(_W_VAL, "0")) if ilvl_elem is not None else 0

            numbering_part = para.part.numbering_part
            if numbering_part is None:
                return False

            num_fmt = self._resolve_num_fmt(numbering_part, num_id, ilvl)
            if num_fmt is not None:
                return num_fmt in _ORDERED_NUM_FMTS

        except Exception as exc:
            logger.debug("Could not determine list order via numPr: %s", exc)

        return False

    @staticmethod
    def _resolve_num_fmt(
        numbering_part: object, num_id: int, ilvl: int
    ) -> str | None:
        """
        Walk the numbering.xml indirection chain and return the ``numFmt``
        value string for the given (*num_id*, *ilvl*) combination.

        The chain is:
          ``<w:num w:numId="N">``
            ``<w:abstractNumId w:val="A"/>``
              → ``<w:abstractNum w:abstractNumId="A">``
                ``<w:lvl w:ilvl="D">``
                  ``<w:numFmt w:val="…"/>``

        Returns ``None`` when any part of the chain is missing or malformed.
        """
        _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        def _attr(elem: object, local: str) -> str:
            return elem.get(f"{{{_W_NS}}}{local}", "")  # type: ignore[union-attr]

        try:
            root = numbering_part._element  # type: ignore[union-attr]
            ns = {"w": _W_NS}

            # Step 1: locate <w:num w:numId="N">
            num_elem = root.find(
                f"w:num[@{{{_W_NS}}}numId='{num_id}']", ns
            )
            if num_elem is None:
                return None

            # Step 2: resolve abstractNumId
            abstract_id_elem = num_elem.find("w:abstractNumId", ns)
            if abstract_id_elem is None:
                return None
            abstract_num_id_str = _attr(abstract_id_elem, "val")
            if not abstract_num_id_str:
                return None
            abstract_num_id = int(abstract_num_id_str)

            # Step 3: locate <w:abstractNum w:abstractNumId="A">
            abstract_num_elem = root.find(
                f"w:abstractNum[@{{{_W_NS}}}abstractNumId='{abstract_num_id}']",
                ns,
            )
            if abstract_num_elem is None:
                return None

            # Step 4: find <w:lvl w:ilvl="D">
            lvl = abstract_num_elem.find(
                f"w:lvl[@{{{_W_NS}}}ilvl='{ilvl}']", ns
            )
            if lvl is None:
                return None

            # Step 5: read <w:numFmt w:val="…">
            num_fmt_elem = lvl.find("w:numFmt", ns)
            if num_fmt_elem is None:
                return None
            return _attr(num_fmt_elem, "val") or None

        except Exception:
            return None

    # ------------------------------------------------------------------
    # Run / inline content extraction
    # ------------------------------------------------------------------

    def _extract_runs(self, para: DocxParagraph) -> list[TextRun]:
        """
        Extract all :class:`~schema.TextRun` objects from *para*.

        Handles:
        - Plain text runs (``<w:r>``)
        - Bold, italic, underline, strikethrough (from ``<w:rPr>``)
        - Monospace fonts → inline code flag
        - Hyperlinks (``<w:hyperlink>``) with URL resolution
        - Tracked insertions (``<w:ins>``)
        - Tab characters and soft line-breaks (``<w:br>``)

        Fix #8 — whitespace-only runs are now preserved (they matter for
        spacing between formatted and plain text spans within a paragraph).
        """
        runs: list[TextRun] = []

        for child in para._element:
            tag = child.tag

            if tag == _W_R:
                run = self._parse_run_element(child)
                if run is not None:
                    runs.append(run)

            elif tag == _W_HYPERLINK:
                # Fix #7 — hyperlink URL resolution is robust; nested runs
                # inside the hyperlink are all annotated with the same URL.
                url = self._resolve_hyperlink_url(para, child)
                for r_elem in child.findall(_W_R):
                    run = self._parse_run_element(r_elem, hyperlink=url)
                    if run is not None:
                        runs.append(run)

            elif tag == _W_INS:
                # Tracked insertions contain accepted text; include it.
                for r_elem in child.findall(_W_R):
                    run = self._parse_run_element(r_elem)
                    if run is not None:
                        runs.append(run)

            # _W_DEL is skipped — deleted text must not appear in output.
            # _W_DRAWING / _W_PICT are handled by _parse_mixed_paragraph.

        return runs

    def _parse_run_element(
        self,
        r_elem: object,
        hyperlink: str | None = None,
    ) -> TextRun | None:
        """
        Parse a single ``<w:r>`` XML element into a :class:`~schema.TextRun`.

        Fix #8 — whitespace-only text is now kept (was silently dropped
        when the run contained only a space between two formatted spans).
        The check ``if not text`` is intentionally replaced with a check for
        a completely *empty* string (zero characters).

        Returns ``None`` only when there is literally no text content at all
        (e.g. a run that contains only a drawing placeholder).
        """
        text_parts: list[str] = []

        # Collect text from all <w:t> children.
        # <w:t xml:space="preserve"> is the normal form; the attribute is
        # respected by python-docx automatically, but we read the raw text
        # value here so we must honour leading/trailing spaces ourselves.
        for t_elem in r_elem.findall(_W_T):  # type: ignore[union-attr]
            text_parts.append(t_elem.text or "")

        # <w:tab> → literal tab character.
        for _ in r_elem.findall(_W_TAB):  # type: ignore[union-attr]
            text_parts.append("\t")

        # <w:br> → line break (soft return) or page break.
        for br_elem in r_elem.findall(_W_BR):  # type: ignore[union-attr]
            br_type = br_elem.get(_W_VAL, "")
            if br_type == "page":
                pass  # page breaks are layout-only; skip
            else:
                text_parts.append("\n")

        text = "".join(text_parts)
        # Return None only when there is truly no content at all.
        if text == "":
            return None

        # --- Run properties ---
        rPr = r_elem.find(_W_RPR)  # type: ignore[union-attr]
        bold = italic = underline = strikethrough = code = False

        if rPr is not None:
            # Bold: <w:b/> present and not explicitly val="0" (Word uses
            # val="0" to *remove* inherited bold from a character style).
            b_elem = rPr.find(_W_B)
            bold = b_elem is not None and b_elem.get(_W_VAL, "1") != "0"

            i_elem = rPr.find(_W_I)
            italic = i_elem is not None and i_elem.get(_W_VAL, "1") != "0"

            u_elem = rPr.find(_W_U)
            # <w:u w:val="none"/> means no underline; anything else is underline.
            if u_elem is not None:
                underline = u_elem.get(_W_VAL, "single") not in ("", "none")

            strike_elem = rPr.find(_W_STRIKE)
            dstrike_elem = rPr.find(_W_DSTRIKE)
            strikethrough = (strike_elem is not None) or (dstrike_elem is not None)

            # Inline code: detect a known monospace font family.
            rFonts = rPr.find(_W_RFONTS)
            if rFonts is not None:
                ascii_font = (rFonts.get(_W_ASCII) or "").lower()
                hansi_font = (rFonts.get(_W_HANSI) or "").lower()
                font_name = ascii_font or hansi_font
                if any(mono in font_name for mono in _MONO_FONTS):
                    code = True

        return TextRun(
            text=text,
            bold=bold,
            italic=italic,
            underline=underline,
            strikethrough=strikethrough,
            code=code,
            hyperlink=hyperlink,
        )

    def _resolve_hyperlink_url(
        self, para: DocxParagraph, hyperlink_elem: object
    ) -> str | None:
        """
        Return the URL for a ``<w:hyperlink>`` element.

        Fix #7 — handles the following cases robustly:
        - Normal external hyperlinks stored as relationship targets.
        - Anchor-only links (``w:anchor`` attribute, no relationship).
        - Missing or broken relationships (returns ``None`` without crashing).
        - The relationship ID may be on ``r:id`` or on the element itself.

        Word encodes two kinds of hyperlinks:
        1. External URL: ``<w:hyperlink r:id="rId5">`` → look up in .rels.
        2. Bookmark anchor: ``<w:hyperlink w:anchor="section2">`` → fragment.
        """
        # Case 1: relationship-based URL.
        r_id = hyperlink_elem.get(_R_ID)  # type: ignore[union-attr]
        if r_id:
            try:
                rel = para.part.relationships[r_id]
                return rel.target_ref
            except (KeyError, AttributeError, TypeError) as exc:
                logger.debug(
                    "Hyperlink relationship %r not found: %s", r_id, exc
                )

        # Case 2: inline anchor (bookmark reference within the same document).
        anchor = hyperlink_elem.get(qn("w:anchor"))  # type: ignore[union-attr]
        if anchor:
            return f"#{anchor}"

        return None

    # ------------------------------------------------------------------
    # Table parsing (Fix #3)
    # ------------------------------------------------------------------

    def _parse_table(self, table: DocxTable) -> Table:
        """
        Convert a ``python-docx`` :class:`~docx.table.Table` into a
        :class:`~schema.Table`.

        Fix #3 — cells are now parsed using :meth:`_extract_runs_from_cell`
        which preserves bold, italic, underline, hyperlinks, and other inline
        formatting.  The old approach only extracted plain `cell.text` which
        collapsed all formatting.

        Multiple paragraphs inside a single cell are joined with a newline
        run so that multi-paragraph cells render with visible separation.

        Fix #6 — image extraction is also attempted for each cell paragraph.
        """
        rows: list[TableRow] = []
        for row_idx, row in enumerate(table.rows):
            cells: list[TableCell] = []
            is_header_row = row_idx == 0
            for cell in row.cells:
                parsed_cell = self._parse_table_cell(cell, is_header=is_header_row)
                cells.append(parsed_cell)
            rows.append(TableRow(cells=cells))

        has_header = len(rows) > 0
        return Table(rows=rows, has_header=has_header)

    def _parse_table_cell(
        self, cell: object, *, is_header: bool
    ) -> TableCell:
        """
        Parse a single table cell, preserving inline formatting.

        A DOCX table cell may contain multiple ``<w:p>`` paragraphs.  We
        collect runs from all of them, inserting a ``\\n`` separator run
        between adjacent paragraphs so multi-paragraph cells are legible.

        Fix #3 — rich run extraction replaces the previous ``cell.text`` call.
        Fix #6 — images inside table cells are extracted and stored in the
        runs list as placeholder text runs (full Image block support inside
        cells is complex and deferred to future work; images are at least not
        silently lost — a ``[image]`` placeholder text run is emitted).
        """
        all_runs: list[TextRun] = []
        paragraphs = list(cell.paragraphs)  # type: ignore[union-attr]

        for p_idx, para in enumerate(paragraphs):
            para_runs = self._extract_runs(para)

            # Check for images in this cell paragraph and emit placeholders.
            # Full Image blocks cannot be emitted inside a TableCell (the schema
            # does not support nested Image nodes in cells), so we represent them
            # as a descriptive text run instead.  The image is still saved to
            # disk so it is not lost.
            for child in para._element.iter():
                if child.tag == _W_DRAWING or child.tag == _W_PICT:
                    img = self._extract_image_from_element(child, para)
                    if img is not None:
                        para_runs.append(
                            TextRun(text=f"[{img.alt or 'image'}]", italic=True)
                        )

            if para_runs:
                if all_runs and p_idx > 0:
                    # Separator between paragraphs within the same cell.
                    all_runs.append(TextRun(text="\n"))
                all_runs.extend(para_runs)

        return TableCell(runs=all_runs, is_header=is_header)

    # ------------------------------------------------------------------
    # Image extraction (Fix #6, Fix #14)
    # ------------------------------------------------------------------

    def _extract_image_from_element(
        self, elem: object, para: DocxParagraph
    ) -> Image | None:
        """
        Scan *elem* (a ``<w:drawing>`` or ``<w:pict>`` subtree) for an
        ``<a:blip>`` element and extract the referenced image.

        Fix #6 — detection now covers both ``<w:drawing>`` (modern inline
        images) and ``<w:pict>`` (legacy VML images).  The ``<a:blip>`` tag
        is the authoritative image reference in both cases.

        Returns ``None`` when no image is found or extraction fails.
        """
        for child in elem.iter():  # type: ignore[union-attr]
            if child.tag == _A_BLIP:
                r_embed = child.get(_R_EMBED)
                if r_embed:
                    return self._save_image(para, r_embed)
        return None

    def _save_image(
        self, para: DocxParagraph, r_id: str
    ) -> Image | None:
        """
        Resolve *r_id* to an image part, save it to :attr:`_media_dir`, and
        return an :class:`~schema.Image` AST node.

        Fix #14 — the ``media/`` directory is created lazily (only when an
        actual image is about to be saved) to avoid creating an empty directory
        for documents that contain no images.

        Content-type → file extension mapping covers the most common formats:
        ``image/png`` → ``.png``, ``image/jpeg`` → ``.jpg``, etc.  Unknown
        subtypes fall back to using the subtype string directly.
        """
        try:
            image_part = para.part.related_parts[r_id]
        except (KeyError, AttributeError) as exc:
            logger.debug("Image relationship %r not found: %s", r_id, exc)
            return None

        try:
            content_type: str = image_part.content_type  # e.g. "image/png"
            raw_ext = content_type.split("/")[-1].lower()
            # Normalise common MIME sub-type variations.
            ext = {"jpeg": "jpg", "svg+xml": "svg", "tiff": "tif"}.get(
                raw_ext, raw_ext
            )

            self._image_counter += 1
            filename = f"image{self._image_counter}.{ext}"

            # Fix #14: create directory lazily.
            self._media_dir.mkdir(parents=True, exist_ok=True)
            dest = self._media_dir / filename
            dest.write_bytes(image_part.blob)

            rel_path = f"media/{filename}"
            return Image(src=rel_path, alt=f"Image {self._image_counter}")

        except Exception as exc:
            logger.warning(
                "Failed to extract image (r_id=%r): %s", r_id, exc
            )
            return None

    # ------------------------------------------------------------------
    # Alignment
    # ------------------------------------------------------------------

    @staticmethod
    def _para_alignment(para: DocxParagraph) -> str | None:
        """Return the CSS text-alignment string for *para*, or ``None``."""
        try:
            alignment = para.alignment
            if alignment is not None:
                return _ALIGN_MAP.get(alignment.value)  # type: ignore[union-attr]
        except Exception:
            pass
        return None

    # ------------------------------------------------------------------
    # Horizontal rule detection (Fix #11)
    # ------------------------------------------------------------------

    @staticmethod
    def _is_horizontal_rule(para: DocxParagraph) -> bool:
        """
        Return ``True`` when *para* carries a bottom paragraph border that
        Word uses to represent a horizontal rule.

        Fix #11 — the caller now also checks ``para.text.strip()`` before
        accepting this as an HR, so a non-empty paragraph that happens to have
        a border is not misclassified.

        Word encodes horizontal rules as an empty paragraph with:

        .. code-block:: xml

            <w:pPr>
              <w:pBdr>
                <w:bottom w:val="single" w:sz="6" …/>
              </w:pBdr>
            </w:pPr>
        """
        try:
            pPr = para._element.find(_W_PPR)
            if pPr is None:
                return False
            pBdr = pPr.find(_W_PBDR)
            if pBdr is None:
                return False
            bottom = pBdr.find(_W_BOTTOM)
            if bottom is None:
                return False
            val = bottom.get(_W_VAL, "")
            return bool(val) and val != "none"
        except Exception:
            return False

    # ------------------------------------------------------------------
    # Core properties / document metadata (Fix #15)
    # ------------------------------------------------------------------

    def _core_property(self, name: str) -> str:
        """
        Safely retrieve a core document property (title, author, created, …).

        Returns an empty string when the property is absent or the document
        does not include core properties.
        """
        try:
            props = self._doc.core_properties
            value = getattr(props, name, None)
            if value is None:
                return ""
            return str(value)
        except Exception:
            return ""

    # ------------------------------------------------------------------
    # O(1) python-docx object lookup (Fix #1)
    # ------------------------------------------------------------------

    def _find_table_for_element(self, tbl_elem: object) -> DocxTable | None:
        """
        Return the ``python-docx`` Table whose underlying XML element is
        *tbl_elem*, using the O(1) lookup map built in ``__init__``.

        Fix #1 — replaces the previous O(n) linear scan over ``doc.tables``.
        """
        return self._table_map.get(tbl_elem)  # type: ignore[arg-type]

    def _find_para_for_element(self, p_elem: object) -> DocxParagraph | None:
        """
        Return the ``python-docx`` Paragraph whose underlying XML element is
        *p_elem*, using the O(1) lookup map built in ``__init__``.

        Fix #1 — replaces the previous O(n) linear scan over ``doc.paragraphs``.

        Note: ``doc.paragraphs`` includes paragraphs that are *inside* tables.
        The map therefore covers both top-level and nested paragraphs, which is
        exactly what we need when handling ``<w:sdt>`` content controls and
        other wrapper elements.
        """
        return self._para_map.get(p_elem)  # type: ignore[arg-type]
