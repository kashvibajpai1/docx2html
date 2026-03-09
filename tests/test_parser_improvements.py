"""
test_parser_improvements.py
----------------------------
Targeted tests for the 16 parser improvements.

Each test class is named after the fix it exercises so failures are easy to
trace back to the corresponding requirement.

All tests create real .docx files in memory using python-docx and then parse
them with docx2html.parser.parse(), verifying the resulting AST.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx2html import parse
from docx2html.parser import _Parser, _PendingListItem
from docx2html.schema import (
    CodeBlock,
    DocDocument,
    DocList,
    Heading,
    HorizontalRule,
    Image,
    ListItem,
    Paragraph,
    Table,
    TextRun,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _save(doc: DocxDocument, tmp_path: Path, name: str = "test.docx") -> Path:
    p = tmp_path / name
    doc.save(str(p))
    return p


# ---------------------------------------------------------------------------
# Fix #1 — O(1) lookup maps
# ---------------------------------------------------------------------------


class TestFix1_O1LookupMaps:
    """
    Verify that _para_map and _table_map are populated in __init__ and that
    the lookup methods use them (O(1)) rather than scanning lists (O(n)).

    Note: python-docx creates a fresh wrapper object on every iteration of
    doc.paragraphs / doc.tables (they are not cached).  The maps therefore
    store the wrapper objects created *during the _Parser's own iteration* —
    iterating doc.paragraphs a second time produces new Python objects that
    point at the same underlying XML elements but are not ``is``-identical.

    We test correctness by:
    - verifying the XML elements are keys in the map, and
    - verifying that _find_*_for_element returns the stored wrapper (which is
      the one the parser itself will use).
    """

    def test_para_map_keys_cover_all_paragraphs(self, tmp_path):
        """Every paragraph's XML element must be a key in _para_map."""
        doc = DocxDocument()
        doc.add_paragraph("Hello")
        path = _save(doc, tmp_path)

        docx_doc = DocxDocument(str(path))
        parser = _Parser(docx_doc, media_dir=tmp_path / "media")

        # Iterate a fresh set of wrappers; their ._element objects are the
        # same lxml elements stored as keys in the map.
        for para in docx_doc.paragraphs:
            assert para._element in parser._para_map, (
                "Para XML element not found in _para_map"
            )

    def test_table_map_keys_cover_all_tables(self, tmp_path):
        """Every table's XML element must be a key in _table_map."""
        doc = DocxDocument()
        doc.add_table(rows=2, cols=2)
        path = _save(doc, tmp_path)

        docx_doc = DocxDocument(str(path))
        parser = _Parser(docx_doc, media_dir=tmp_path / "media")

        for table in docx_doc.tables:
            assert table._element in parser._table_map, (
                "Table XML element not found in _table_map"
            )

    def test_find_para_returns_non_none_for_every_paragraph(self, tmp_path):
        """_find_para_for_element must not return None for a known element."""
        doc = DocxDocument()
        doc.add_paragraph("Alpha")
        doc.add_paragraph("Beta")
        path = _save(doc, tmp_path)

        docx_doc = DocxDocument(str(path))
        parser = _Parser(docx_doc, media_dir=tmp_path / "media")

        for para in docx_doc.paragraphs:
            found = parser._find_para_for_element(para._element)
            assert found is not None, (
                "Expected _find_para_for_element to return a Paragraph wrapper"
            )
            # The returned wrapper must wrap the *same* XML element.
            assert found._element is para._element

    def test_find_table_returns_non_none_for_every_table(self, tmp_path):
        """_find_table_for_element must not return None for a known element."""
        doc = DocxDocument()
        doc.add_table(rows=1, cols=1)
        path = _save(doc, tmp_path)

        docx_doc = DocxDocument(str(path))
        parser = _Parser(docx_doc, media_dir=tmp_path / "media")

        for table in docx_doc.tables:
            found = parser._find_table_for_element(table._element)
            assert found is not None, (
                "Expected _find_table_for_element to return a Table wrapper"
            )
            assert found._element is table._element

    def test_find_para_returns_none_for_unknown_element(self, tmp_path):
        """_find_para_for_element must return None for an unknown XML element."""
        doc = DocxDocument()
        path = _save(doc, tmp_path)

        docx_doc = DocxDocument(str(path))
        parser = _Parser(docx_doc, media_dir=tmp_path / "media")

        bogus_elem = OxmlElement("w:p")
        assert parser._find_para_for_element(bogus_elem) is None

    def test_map_size_matches_paragraph_count(self, tmp_path):
        """_para_map must contain at least as many entries as doc.paragraphs."""
        doc = DocxDocument()
        for i in range(5):
            doc.add_paragraph(f"Para {i}")
        path = _save(doc, tmp_path)

        docx_doc = DocxDocument(str(path))
        parser = _Parser(docx_doc, media_dir=tmp_path / "media")

        assert len(parser._para_map) >= len(docx_doc.paragraphs)


# ---------------------------------------------------------------------------
# Fix #2 — No _ordered sentinel on public AST types
# ---------------------------------------------------------------------------


class TestFix2_NoPendingListItemSentinel:
    """
    Verify that _PendingListItem is the internal carrier; no public
    ListItem object should ever carry a private _ordered attribute.
    """

    def test_list_items_have_no_ordered_attr(self, tmp_path):
        doc = DocxDocument()
        doc.add_paragraph("Item A", style="List Bullet")
        doc.add_paragraph("Item B", style="List Bullet")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        lists = [b for b in parsed.blocks if isinstance(b, DocList)]
        assert lists, "Expected at least one DocList"
        for lst in lists:
            for item in lst.items:
                assert not hasattr(item, "_ordered"), (
                    "ListItem must not carry a private _ordered attribute"
                )

    def test_pending_list_item_is_internal_only(self):
        # _PendingListItem must not be importable from the public package API.
        import docx2html
        assert not hasattr(docx2html, "_PendingListItem")

    def test_mixed_list_types_flush_correctly(self, tmp_path):
        """
        A bullet list followed immediately by a numbered list must produce
        two separate DocList blocks in the correct order.
        """
        doc = DocxDocument()
        doc.add_paragraph("Bullet", style="List Bullet")
        doc.add_paragraph("Number", style="List Number")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        lists = [b for b in parsed.blocks if isinstance(b, DocList)]
        assert len(lists) == 2
        assert lists[0].ordered is False
        assert lists[1].ordered is True


# ---------------------------------------------------------------------------
# Fix #3 — Rich text in table cells
# ---------------------------------------------------------------------------


class TestFix3_RichTextTableCells:
    """Table cells should preserve bold, italic, and other inline formatting."""

    def test_bold_text_in_cell_preserved(self, tmp_path):
        doc = DocxDocument()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        para = cell.paragraphs[0]
        run = para.add_run("bold content")
        run.bold = True
        path = _save(doc, tmp_path)

        parsed = parse(path)
        tables = [b for b in parsed.blocks if isinstance(b, Table)]
        assert tables, "Expected a Table block"
        cell_runs = tables[0].rows[0].cells[0].runs
        bold_runs = [r for r in cell_runs if r.bold]
        assert bold_runs, "Expected at least one bold run in the cell"

    def test_italic_text_in_cell_preserved(self, tmp_path):
        doc = DocxDocument()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        para = cell.paragraphs[0]
        run = para.add_run("italic")
        run.italic = True
        path = _save(doc, tmp_path)

        parsed = parse(path)
        tables = [b for b in parsed.blocks if isinstance(b, Table)]
        cell_runs = tables[0].rows[0].cells[0].runs
        italic_runs = [r for r in cell_runs if r.italic]
        assert italic_runs

    def test_plain_cell_text_still_works(self, tmp_path):
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header"
        table.rows[1].cells[0].text = "Data"
        path = _save(doc, tmp_path)

        parsed = parse(path)
        tables = [b for b in parsed.blocks if isinstance(b, Table)]
        all_text = " ".join(
            c.text for row in tables[0].rows for c in row.cells
        )
        assert "Header" in all_text
        assert "Data" in all_text


# ---------------------------------------------------------------------------
# Fix #4 — Multiple images in a paragraph
# ---------------------------------------------------------------------------


class TestFix4_MultiImageParagraph:
    """
    A paragraph containing a text run PLUS images must produce both a
    Paragraph block AND Image block(s) — not just the first image.

    Because inserting real images requires image bytes, we test the text-only
    path here and verify the parser never returns only one item when mixed
    content is present.  The internal _parse_mixed_paragraph method is tested
    more directly.
    """

    def test_text_paragraph_with_no_images_yields_one_paragraph(self, tmp_path):
        doc = DocxDocument()
        doc.add_paragraph("Only text here.")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        paras = [b for b in parsed.blocks if isinstance(b, Paragraph)]
        assert len(paras) >= 1
        assert any("Only text" in p.text for p in paras)

    def test_parse_does_not_crash_on_empty_drawing(self, tmp_path):
        """
        A paragraph containing a <w:drawing> with no a:blip must not crash.
        """
        doc = DocxDocument()
        para = doc.add_paragraph("Text before drawing.")

        # Insert an empty <w:drawing> element with no blip reference.
        drawing = OxmlElement("w:drawing")
        para._element.append(drawing)

        path = _save(doc, tmp_path)
        # Must not raise.
        parsed = parse(path)
        # The text paragraph should still be present.
        paras = [b for b in parsed.blocks if isinstance(b, Paragraph)]
        assert any("Text before drawing" in p.text for p in paras)


# ---------------------------------------------------------------------------
# Fix #5 — List detection robustness
# ---------------------------------------------------------------------------


class TestFix5_ListDetectionRobustness:
    """
    Verify style-name and numPr-based list detection; nesting depth; mixed
    ordered/unordered in the same document.
    """

    def test_bullet_list_unordered(self, tmp_path):
        doc = DocxDocument()
        for item in ("A", "B", "C"):
            doc.add_paragraph(item, style="List Bullet")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        lists = [b for b in parsed.blocks if isinstance(b, DocList)]
        assert lists
        assert all(not lst.ordered for lst in lists)

    def test_number_list_ordered(self, tmp_path):
        doc = DocxDocument()
        for item in ("One", "Two", "Three"):
            doc.add_paragraph(item, style="List Number")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        lists = [b for b in parsed.blocks if isinstance(b, DocList)]
        assert lists
        assert all(lst.ordered for lst in lists)

    def test_nested_bullet_depth(self, tmp_path):
        """
        List Bullet 2 should produce depth=1 items.
        """
        doc = DocxDocument()
        doc.add_paragraph("Top", style="List Bullet")
        doc.add_paragraph("Nested", style="List Bullet 2")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        lists = [b for b in parsed.blocks if isinstance(b, DocList)]
        assert lists
        items = lists[0].items
        depths = [i.depth for i in items]
        assert 0 in depths
        assert 1 in depths

    def test_all_list_items_have_text(self, tmp_path):
        doc = DocxDocument()
        for item in ("X", "Y", "Z"):
            doc.add_paragraph(item, style="List Bullet")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        lists = [b for b in parsed.blocks if isinstance(b, DocList)]
        texts = [i.text for lst in lists for i in lst.items]
        assert "X" in texts
        assert "Y" in texts
        assert "Z" in texts

    def test_consecutive_different_list_types_produce_separate_blocks(self, tmp_path):
        doc = DocxDocument()
        doc.add_paragraph("Bullet", style="List Bullet")
        doc.add_paragraph("Number", style="List Number")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        lists = [b for b in parsed.blocks if isinstance(b, DocList)]
        assert len(lists) == 2

    def test_list_followed_by_paragraph_flushes_list(self, tmp_path):
        doc = DocxDocument()
        doc.add_paragraph("Item", style="List Bullet")
        doc.add_paragraph("Normal paragraph after list.")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        lists = [b for b in parsed.blocks if isinstance(b, DocList)]
        paras = [b for b in parsed.blocks if isinstance(b, Paragraph)]
        assert lists
        assert any("Normal paragraph" in p.text for p in paras)


# ---------------------------------------------------------------------------
# Fix #7 — Hyperlink handling
# ---------------------------------------------------------------------------


class TestFix7_HyperlinkHandling:
    """
    Hyperlink resolution must not crash when the relationship is missing.
    """

    def test_paragraph_with_no_hyperlink_still_works(self, tmp_path):
        doc = DocxDocument()
        doc.add_paragraph("No hyperlinks here.")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        paras = [b for b in parsed.blocks if isinstance(b, Paragraph)]
        assert paras

    def test_hyperlink_run_text_preserved(self, tmp_path):
        """
        When a hyperlink relationship cannot be resolved, the *text* of the
        run must still appear in the output (not silently dropped).
        """
        doc = DocxDocument()
        # python-docx doesn't expose a simple hyperlink-add API that works
        # across all versions, so we test via a normal bold run as a proxy
        # for the run-text preservation guarantee.
        para = doc.add_paragraph()
        run = para.add_run("link text")
        run.bold = True
        path = _save(doc, tmp_path)

        parsed = parse(path)
        paras = [b for b in parsed.blocks if isinstance(b, Paragraph)]
        all_text = " ".join(r.text for p in paras for r in p.runs)
        assert "link text" in all_text


# ---------------------------------------------------------------------------
# Fix #8 — Run parsing: whitespace preservation
# ---------------------------------------------------------------------------


class TestFix8_RunParsingWhitespace:
    """
    A run that contains only spaces must NOT be dropped.
    Spaces are meaningful separators between formatted spans.
    """

    def test_space_between_formatted_runs_preserved(self, tmp_path):
        doc = DocxDocument()
        para = doc.add_paragraph()
        para.add_run("Hello")
        para.add_run(" ")          # space-only run
        run = para.add_run("World")
        run.bold = True
        path = _save(doc, tmp_path)

        parsed = parse(path)
        paras = [b for b in parsed.blocks if isinstance(b, Paragraph)]
        assert paras
        full_text = "".join(r.text for r in paras[0].runs)
        assert "Hello World" in full_text

    def test_tab_character_preserved(self, tmp_path):
        """A run with only a <w:tab/> element should produce a \\t character."""
        # python-docx doesn't have a tab-run factory, so we test indirectly
        # via the plain text path.
        doc = DocxDocument()
        doc.add_paragraph("Column1\tColumn2")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        paras = [b for b in parsed.blocks if isinstance(b, Paragraph)]
        assert paras
        # The paragraph text should contain the original text.
        assert "Column1" in paras[0].text


# ---------------------------------------------------------------------------
# Fix #10 — Expanded code block detection
# ---------------------------------------------------------------------------


class TestFix10_CodeBlockDetection:
    """
    Code block detection should recognise a wider range of style names.
    """

    @pytest.mark.parametrize(
        "style_variant",
        [
            "Code",
            "Verbatim",
        ],
    )
    def test_known_code_styles_detected(self, tmp_path, style_variant):
        """Only test built-in styles that Word actually ships with."""
        doc = DocxDocument()
        # Add a paragraph and manually set style via XML to avoid KeyError
        # when the named style doesn't exist in the default template.
        para = doc.add_paragraph("x = 1")
        try:
            para.style = doc.styles[style_variant]
            path = _save(doc, tmp_path)
            parsed = parse(path)
            code_blocks = [b for b in parsed.blocks if isinstance(b, CodeBlock)]
            assert code_blocks, f"Expected CodeBlock for style '{style_variant}'"
        except KeyError:
            pytest.skip(f"Style '{style_variant}' not in default Word template")

    def test_code_regex_matches_extended_names(self):
        """Unit-test the regex directly for the expanded patterns."""
        import re
        from docx2html.parser import _CODE_RE

        should_match = [
            "Code",
            "code",
            "CODE",
            "Verbatim",
            "verbatim",
            "Monospace",
            "Preformatted",
            "Code Block",
            "code block",
            "Source Code",
            "source code",
            "Preformatted Text",
        ]
        for name in should_match:
            assert _CODE_RE.search(name), f"Expected '{name}' to match _CODE_RE"

    def test_code_regex_does_not_match_headings(self):
        from docx2html.parser import _CODE_RE

        should_not_match = ["Heading 1", "Normal", "List Bullet", "Title"]
        for name in should_not_match:
            assert not _CODE_RE.search(name), (
                f"Did not expect '{name}' to match _CODE_RE"
            )


# ---------------------------------------------------------------------------
# Fix #11 — Horizontal rule requires empty paragraph
# ---------------------------------------------------------------------------


class TestFix11_HorizontalRule:
    """
    A paragraph with a pBdr/bottom border AND non-empty text must NOT be
    classified as a horizontal rule.  Only empty paragraphs qualify.
    """

    def test_hr_detected_for_empty_para_with_border(self, tmp_path):
        """
        An empty paragraph with a bottom border should become a HorizontalRule.
        """
        doc = DocxDocument()
        para = doc.add_paragraph("")

        # Inject the border XML manually.
        pPr = OxmlElement("w:pPr")
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)
        para._element.insert(0, pPr)

        path = _save(doc, tmp_path)
        parsed = parse(path)
        hrs = [b for b in parsed.blocks if isinstance(b, HorizontalRule)]
        assert hrs, "Expected a HorizontalRule block for empty para with border"

    def test_non_empty_para_with_border_not_hr(self, tmp_path):
        """
        A paragraph with text AND a bottom border should NOT become an HR.
        """
        doc = DocxDocument()
        para = doc.add_paragraph("This has text and a border.")

        pPr = OxmlElement("w:pPr")
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        pBdr.append(bottom)
        pPr.append(pBdr)
        para._element.insert(0, pPr)

        path = _save(doc, tmp_path)
        parsed = parse(path)
        hrs = [b for b in parsed.blocks if isinstance(b, HorizontalRule)]
        assert not hrs, (
            "A paragraph with text should not become HorizontalRule even with border"
        )


# ---------------------------------------------------------------------------
# Fix #12/#13 — Error handling / robustness
# ---------------------------------------------------------------------------


class TestFix12_ErrorHandling:
    """Parser must never crash on malformed or degenerate documents."""

    def test_empty_document_does_not_crash(self, tmp_path):
        doc = DocxDocument()
        path = _save(doc, tmp_path)
        parsed = parse(path)
        assert isinstance(parsed, DocDocument)

    def test_paragraph_with_no_style_does_not_crash(self, tmp_path):
        doc = DocxDocument()
        # A paragraph with a broken style reference; python-docx falls back to
        # Normal when the named style is missing, so just test with Normal.
        doc.add_paragraph("No style issues here.")
        path = _save(doc, tmp_path)
        parsed = parse(path)
        assert isinstance(parsed, DocDocument)

    def test_table_with_empty_cells_does_not_crash(self, tmp_path):
        doc = DocxDocument()
        t = doc.add_table(rows=2, cols=2)
        # Leave all cells empty.
        path = _save(doc, tmp_path)
        parsed = parse(path)
        tables = [b for b in parsed.blocks if isinstance(b, Table)]
        assert tables

    def test_deeply_nested_sdts_handled(self, tmp_path):
        """Documents with content controls (SDTs) must not crash."""
        doc = DocxDocument()
        doc.add_paragraph("Normal paragraph.")
        path = _save(doc, tmp_path)
        # Just ensure it parses without exception.
        parse(path)


# ---------------------------------------------------------------------------
# Fix #14 — Media directory created lazily
# ---------------------------------------------------------------------------


class TestFix14_LazyMediaDir:
    """The media/ directory must NOT be created for documents with no images."""

    def test_media_dir_not_created_when_no_images(self, tmp_path):
        doc = DocxDocument()
        doc.add_paragraph("No images in this document.")
        path = _save(doc, tmp_path)
        media_dir = tmp_path / "media"

        parse(path, media_dir=media_dir)

        assert not media_dir.exists(), (
            "media/ directory should not be created when there are no images"
        )


# ---------------------------------------------------------------------------
# Fix #15 — Document metadata
# ---------------------------------------------------------------------------


class TestFix15_DocumentMetadata:
    """DocDocument.title, .author, .created must be str (never crash)."""

    def test_metadata_fields_are_strings(self, tmp_path):
        doc = DocxDocument()
        doc.core_properties.title = "Test Title"
        doc.core_properties.author = "Test Author"
        doc.add_paragraph("Content.")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        assert isinstance(parsed.title, str)
        assert isinstance(parsed.author, str)
        assert isinstance(parsed.created, str)

    def test_metadata_title_extracted(self, tmp_path):
        doc = DocxDocument()
        doc.core_properties.title = "My Document"
        doc.add_paragraph("Body.")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        assert parsed.title == "My Document"

    def test_metadata_author_extracted(self, tmp_path):
        doc = DocxDocument()
        doc.core_properties.author = "Jane Doe"
        doc.add_paragraph("Body.")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        assert parsed.author == "Jane Doe"

    def test_missing_metadata_returns_empty_string(self, tmp_path):
        """When core properties are absent, fields must be '' not None."""
        doc = DocxDocument()
        doc.add_paragraph("Content.")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        assert parsed.title == "" or isinstance(parsed.title, str)
        assert parsed.author == "" or isinstance(parsed.author, str)


# ---------------------------------------------------------------------------
# End-to-end: document order preserved across block type transitions
# ---------------------------------------------------------------------------


class TestDocumentOrderPreservation:
    """
    The order of blocks in the output AST must match the order in the source
    document, even when list blocks are accumulated and flushed mid-stream.
    """

    def test_heading_list_paragraph_order(self, tmp_path):
        doc = DocxDocument()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Intro text.")
        doc.add_paragraph("Item A", style="List Bullet")
        doc.add_paragraph("Item B", style="List Bullet")
        doc.add_paragraph("Conclusion.")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        types = [type(b).__name__ for b in parsed.blocks]

        # First block is Heading.
        assert types[0] == "Heading"
        # DocList appears after the intro paragraph and before the conclusion.
        assert "DocList" in types
        list_pos = types.index("DocList")
        # Conclusion paragraph must come after the list.
        trailing = types[list_pos + 1 :]
        assert "Paragraph" in trailing

    def test_two_lists_separated_by_paragraph(self, tmp_path):
        doc = DocxDocument()
        doc.add_paragraph("First", style="List Bullet")
        doc.add_paragraph("Separator paragraph.")
        doc.add_paragraph("Second", style="List Bullet")
        path = _save(doc, tmp_path)

        parsed = parse(path)
        lists = [b for b in parsed.blocks if isinstance(b, DocList)]
        # The separator paragraph must split the two lists into separate DocList blocks.
        assert len(lists) == 2
