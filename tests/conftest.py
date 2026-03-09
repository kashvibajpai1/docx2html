"""
conftest.py
-----------
Shared pytest fixtures and helpers for the docx2html test suite.

Fixtures here create in-memory .docx documents using python-docx so that
tests run without needing any files on disk.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Inches


# ---------------------------------------------------------------------------
# Helper: write a DocxDocument to a temporary file and return the path
# ---------------------------------------------------------------------------


def _save_docx(doc: DocxDocument, tmp_path: Path, name: str = "test.docx") -> Path:
    path = tmp_path / name
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def docx_with_headings(tmp_path: Path) -> Path:
    """A .docx containing headings at levels 1, 2, and 3."""
    doc = DocxDocument()
    doc.add_heading("Main Title", level=1)
    doc.add_heading("Section One", level=2)
    doc.add_heading("Subsection", level=3)
    doc.add_paragraph("Some body text after the headings.")
    return _save_docx(doc, tmp_path, "headings.docx")


@pytest.fixture()
def docx_with_paragraphs(tmp_path: Path) -> Path:
    """A .docx containing plain paragraphs with various inline styles."""
    doc = DocxDocument()
    p = doc.add_paragraph()
    run = p.add_run("Hello, ")
    run2 = p.add_run("bold world")
    run2.bold = True
    run3 = p.add_run(" and ")
    run4 = p.add_run("italic text")
    run4.italic = True
    run5 = p.add_run(".")
    doc.add_paragraph("A plain second paragraph.")
    return _save_docx(doc, tmp_path, "paragraphs.docx")


@pytest.fixture()
def docx_with_bullet_list(tmp_path: Path) -> Path:
    """A .docx containing an unordered (bullet) list."""
    doc = DocxDocument()
    doc.add_heading("Bullet List", level=1)
    for item in ("Alpha", "Beta", "Gamma"):
        doc.add_paragraph(item, style="List Bullet")
    return _save_docx(doc, tmp_path, "bullets.docx")


@pytest.fixture()
def docx_with_numbered_list(tmp_path: Path) -> Path:
    """A .docx containing an ordered (numbered) list."""
    doc = DocxDocument()
    doc.add_heading("Numbered List", level=1)
    for item in ("First", "Second", "Third"):
        doc.add_paragraph(item, style="List Number")
    return _save_docx(doc, tmp_path, "numbered.docx")


@pytest.fixture()
def docx_with_table(tmp_path: Path) -> Path:
    """A .docx containing a simple 3×3 table."""
    doc = DocxDocument()
    doc.add_heading("Table Example", level=1)
    table = doc.add_table(rows=3, cols=3)
    headers = ["Name", "Age", "City"]
    for col_idx, header in enumerate(headers):
        table.rows[0].cells[col_idx].text = header
    data = [
        ("Alice", "30", "London"),
        ("Bob", "25", "Berlin"),
    ]
    for row_idx, (name, age, city) in enumerate(data, start=1):
        table.rows[row_idx].cells[0].text = name
        table.rows[row_idx].cells[1].text = age
        table.rows[row_idx].cells[2].text = city
    return _save_docx(doc, tmp_path, "table.docx")


@pytest.fixture()
def docx_empty(tmp_path: Path) -> Path:
    """A .docx with no meaningful content (just the default blank document)."""
    doc = DocxDocument()
    return _save_docx(doc, tmp_path, "empty.docx")


@pytest.fixture()
def docx_mixed(tmp_path: Path) -> Path:
    """
    A .docx that exercises many block types in sequence:
    heading → paragraph → bullet list → numbered list → table.
    """
    doc = DocxDocument()
    doc.add_heading("Mixed Document", level=1)
    doc.add_paragraph("Introduction paragraph.")
    doc.add_heading("List Section", level=2)
    for item in ("Bullet A", "Bullet B"):
        doc.add_paragraph(item, style="List Bullet")
    for item in ("Step 1", "Step 2"):
        doc.add_paragraph(item, style="List Number")
    doc.add_heading("Data Table", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Key"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "foo"
    table.rows[1].cells[1].text = "bar"
    return _save_docx(doc, tmp_path, "mixed.docx")
